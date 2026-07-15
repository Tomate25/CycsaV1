import docx

full_path = 'C:/Users/abdia/Downloads/FLUJO COTIZACIONES CYCSA  16062026.docx'
try:
    doc = docx.Document(full_path)
    print("=== FULL CONTENTS OF COTIZACIONES WORKFLOW ===")
    text = []
    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text.strip())
    for t in doc.tables:
        for r in t.rows:
            row_text = [c.text.strip() for c in r.cells if c.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))
    print("\n".join(text))
except Exception as e:
    print("Error:", e)
